"""
Conversión de Markdown a .docx.

Soporta encabezados, listas, tablas, negrita/cursiva e hipervínculos — el
subconjunto que genera el modelo en sus respuestas.
"""
import io
import re
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HEADER_TEXT = "DOCUMENTO CONFIDENCIAL — GESTIÓN DE ASUNTOS PÚBLICOS"
GREY = RGBColor(0x88, 0x88, 0x88)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "1a1a1a"

# Divide un párrafo en tokens de link / negrita / cursiva conservando el resto.
INLINE_RE = re.compile(r'(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|\*[^*]+\*)')
LINK_RE = re.compile(r'\[([^\]]+)\]\(([^)]+)\)')
HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"


def _strip_inline(text: str) -> str:
    """Quita el marcado inline — para celdas de tabla, que van en texto plano."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    return re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)


def _add_hyperlink(paragraph, text: str, url: str):
    """Añade un run con hipervínculo real (python-docx no lo soporta nativamente)."""
    r_id = paragraph.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    r_pr = OxmlElement('w:rPr')
    r_style = OxmlElement('w:rStyle')
    r_style.set(qn('w:val'), 'Hyperlink')
    r_pr.append(r_style)
    run.append(r_pr)

    t = OxmlElement('w:t')
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _fill_inline(paragraph, text: str):
    """Escribe `text` en el párrafo respetando links, negrita y cursiva."""
    for tok in INLINE_RE.split(text):
        link = LINK_RE.match(tok)
        if link:
            _add_hyperlink(paragraph, link.group(1), link.group(2))
        elif tok.startswith('**') and tok.endswith('**'):
            paragraph.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            paragraph.add_run(tok[1:-1]).italic = True
        else:
            paragraph.add_run(tok)


def _is_table_row(s: str) -> bool:
    return s.startswith('|') and s.endswith('|')


def _is_separator_row(s: str) -> bool:
    return _is_table_row(s) and bool(re.match(r'^\|[\s\-|:]+\|$', s))


def _parse_row(s: str) -> list:
    return s.strip('|').split('|')


def _add_table(doc, rows: list):
    tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
    tbl.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = _strip_inline(cell_text.strip())
            if r_idx != 0:
                continue
            # Fila de encabezado: negrita, texto blanco, fondo oscuro.
            para = cell.paragraphs[0]
            para.clear()
            run = para.add_run(_strip_inline(cell_text.strip()))
            run.bold = True
            run.font.color.rgb = WHITE
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), HEADER_FILL)
            cell._tc.get_or_add_tcPr().append(shd)
    doc.add_paragraph()


def _add_banner(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER


def markdown_to_docx(md: str) -> bytes:
    """Convierte Markdown a un .docx y devuelve sus bytes."""
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    _add_banner(doc, HEADER_TEXT)
    doc.add_paragraph()

    lines = md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Bloque de tabla: encabezado + separador + filas.
        if (_is_table_row(line) and i + 1 < len(lines)
                and _is_separator_row(lines[i + 1].rstrip())):
            rows = [_parse_row(line)]
            i += 2  # saltar encabezado y separador
            while i < len(lines) and _is_table_row(lines[i].rstrip()):
                rows.append(_parse_row(lines[i].rstrip()))
                i += 1
            _add_table(doc, rows)
            continue

        if line.startswith('### '):
            doc.add_heading(_strip_inline(line[4:]), level=3)
        elif line.startswith('## '):
            doc.add_heading(_strip_inline(line[3:]), level=2)
        elif line.startswith('# '):
            doc.add_heading(_strip_inline(line[2:]), level=1)
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif re.match(r'^[-*]\s+', line):
            para = doc.add_paragraph(style='List Bullet')
            _fill_inline(para, re.sub(r'^[-*]\s+', '', line))
        elif re.match(r'^\d+\.\s+', line):
            doc.add_paragraph(re.sub(r'^\d+\.\s+', '', line), style='List Number')
        elif line == '':
            doc.add_paragraph()
        else:
            _fill_inline(doc.add_paragraph(), line)
        i += 1

    doc.add_paragraph()
    _add_banner(
        doc,
        f"Generado por Lex — Sistema de Monitoreo Parlamentario · "
        f"{datetime.now().strftime('%d/%m/%Y')}",
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_filename() -> str:
    return f"Resumen-Congreso-{datetime.now().strftime('%Y-%m-%d')}.docx"
