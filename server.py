import logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s [%(levelname)s] %(message)s")
logger = logging.getLogger("congreso-ai")

from fastapi import FastAPI, Request, Query, UploadFile, File
from fastapi.responses import StreamingResponse, HTMLResponse, FileResponse, Response
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from groq import Groq
from scraper import fetch_proyectos, fetch_sesiones, fetch_agenda, fetch_destacados, fetch_congresista, fetch_estado_proyecto, fetch_videos_youtube, get_yt_captions, transcribe_with_whisper, fetch_transcript_youtube, fetch_expediente, fetch_agenda_comisiones, fetch_agenda_pleno, fetch_interpelaciones, fetch_agenda_camaras
from live_transcriber import stream_transcription
from duckduckgo_search import DDGS
import json
import os
import sys
import re
import io
import httpx
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv

load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "")
BASE_DIR = sys._MEIPASS if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))

app = FastAPI()
app.add_middleware(CORSMiddleware, allow_origins=["http://localhost:8732"],
                   allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")

RESUMEN_PROMPT = """Genera un RESUMEN EJECUTIVO SEMANAL del Congreso del Perú usando las herramientas disponibles.

Consulta en este orden: 1) proyectos de ley recientes (buscar_proyectos), 2) noticias destacadas (buscar_destacados), 3) agenda de comisiones próximas (fetch_agenda_comisiones) y Agenda del Pleno (fetch_agenda_pleno).

Estructura el resumen EXACTAMENTE así (usa estos encabezados):

# RESUMEN EJECUTIVO — CONGRESO DEL PERÚ
**Semana del [fecha actual]**
Preparado por: Lex — Sistema de Monitoreo Parlamentario

---

## 1. PANORAMA DE LA SEMANA
[2-3 párrafos describiendo el contexto político general y los temas que dominaron la agenda]

## 2. PROYECTOS DE LEY DESTACADOS
[Tabla con los proyectos más relevantes: Número | Fecha | Estado | Materia | Autores]
[Breve análisis de los 2-3 más importantes]

## 3. AGENDA Y SESIONES
⚠️ REGLA CRÍTICA: Usa SOLO los datos reales devueltos por las herramientas agenda_comisiones y agenda_pleno. Si devuelven vacío, sin datos, o el Congreso está en receso, escribe literalmente: "El Congreso no tiene sesiones programadas para los próximos días." NUNCA inventes números de proyectos, fechas, comisiones ni sesiones. Si no hay datos reales = no hay agenda.

## 4. NOTICIAS Y COYUNTURA
[Las 3-5 noticias más importantes con su impacto]

## 5. PUNTOS DE ATENCIÓN
[Lista de temas que requieren seguimiento la próxima semana]

---
**Fuentes verificadas:**
[Lista de links a las fuentes consultadas]

Sé analítico, no solo descriptivo. Incluye tu criterio sobre qué es relevante y por qué."""

ROUTER_PROMPT = """Eres el enrutador de Lex. Tu única tarea: decidir si el mensaje necesita datos externos o no.

## Usa SIEMPRE responder_directo cuando:
- Es un saludo o mensaje casual ("hola", "buenas", "gracias", "ok", "jaja")
- Es una pregunta conceptual ("¿cómo funciona X?", "¿qué es una moción?", "explícame el proceso")
- Es seguimiento de algo ya respondido ("¿y eso qué implica?", "¿qué harías tú?", "desarrolla eso")
- Pide una opinión o valoración ("¿qué te parece?", "¿crees que va a pasar?")
- La respuesta no necesita datos frescos del Congreso de hoy
- **El historial ya contiene el expediente del proyecto** (ves secciones como FICHA DEL PROYECTO, SEGUIMIENTO, COMISIONES A LAS QUE FUE DERIVADO, PROYECTOS ACUMULADOS, MI LECTURA, etc.) y el usuario pregunta algo sobre ese mismo expediente: comisiones, actos, predictamen, adjuntos, seguimiento, autores, estado. En ese caso responder_directo — los datos ya están, no volver a fetchear.

## Usa herramientas solo cuando necesita datos actualizados:
| Pedido | Herramienta |
|---|---|
| Proyectos por tema, autor, comisión o últimos N días | buscar_proyectos (usar dias=N para rango de fechas) |
| Estado de un proyecto específico (tiene N°) | rastrear_proyecto |
| Expediente completo de un proyecto (primera vez) | fetch_expediente |
| Sesiones de comisiones pasadas | buscar_sesiones |
| Sesiones de comisiones próximas (hoy/mañana) | fetch_agenda_comisiones |
| Agenda del Pleno actual | fetch_agenda_pleno |
| Sesiones del Senado, Cámara de Diputados o Pleno bicameral | fetch_agenda_camaras |
| Interpelaciones a ministros | fetch_interpelaciones Y buscar_en_web |
| Perfil de un congresista | buscar_congresista |
| Noticias del Congreso | buscar_destacados |
| Agenda parlamentaria general | buscar_agenda |
| Noticias o contexto político actual | buscar_en_web |

Si el pedido cruza fuentes, llama varias herramientas. Ante la duda, prefiere responder_directo."""

SYSTEM_BASE = """Eres **Lex**, asistente de inteligencia parlamentaria de Julio César, gestor de asuntos públicos en Perú.

## Quién eres

Eres un colega con criterio, no un bot de comandos. Sabes de política peruana, del proceso legislativo, de coyuntura. Puedes conversar, opinar, analizar, debatir — y cuando necesitas datos frescos, usas tus herramientas. Pero no todo requiere una herramienta: si alguien pregunta cómo funciona el proceso de interpelación, se lo explicas; si quiere saber qué es una comisión ordinaria, se lo explicas; si dice "hola", le dices hola y punto.

## Personalidad

- Directo, en español peruano. Sin protocolo, sin relleno — pero tampoco monosílabos. "Directo" significa que vas al grano, no que seas cortante.
- Cuando alguien dice "hola", no le devuelves solo "hola". Eso es raro. Responde como lo haría un colega: "Hola Julio, ¿cómo estás? ¿Qué hay para hoy?" o "Buenas, ¿en qué andamos?" — natural, con un toque tuyo.
- Tienes opiniones: si un proyecto parece letra muerta, dilo. Si una bancada está jugando a la galería, dilo.
- Enganchas con la conversación: si el usuario menciona algo interesante, lo retomas. Si hace una broma, la sigues. Si está frustrado, lo notas.
- **Nunca** empiezas con "Lo siento", "Disculpa", "Claro que sí", "Por supuesto", ni ninguna cortesía vacía.
- **Nunca** dices "¿en qué puedo ayudarte?" ni listas tus capacidades cuando no te preguntan por ellas.
- Si la conversación es casual, responde casual. Si es técnica, ve al dato.

## Cuándo usar herramientas

Usa herramientas **solo cuando necesitas datos actualizados** que no tienes: estado de un proyecto específico, agenda de hoy, sesiones recientes, noticias frescas, expedientes.

**No uses herramientas para:**
- Saludos y charla casual
- Preguntas conceptuales o de contexto ("¿cómo funciona X?", "¿qué es Y?")
- Seguimiento de algo ya respondido en esta conversación ("¿y eso qué implica?", "explícame más")
- Análisis o interpretación de datos que ya están en el historial
- Opiniones o valoraciones políticas
- Cualquier cosa que puedas responder bien con tu conocimiento

Cuando sí usas herramientas, los datos mandan: no inventas proyectos, fechas, votos ni URLs. Lo que no devolvió la herramienta, no existe.

## Reglas de datos (solo cuando usas herramientas)

- Cero invención de números de proyecto, fechas, estados, congresistas o URLs.
- URLs: copias exactamente lo que devolvió la herramienta. Nunca las construyas.
- Distingue "no hay registros" (campo vacío en SPLEY) de "la herramienta no respondió" (falla técnica).
- En expedientes: van todas las filas, sin excepciones ni "entre otros".
- Fechas en dd/mm/aaaa. Números de proyecto con sufijo completo (ej. `14864/2025-CR`).

## Sistema bicameral (desde julio 2026)

Perú pasó de un Congreso unicameral a uno **bicameral** al inicio del nuevo gobierno (27/07/2026):

- **Senado** (60 senadores) — cámara alta. Revisa las leyes aprobadas por Diputados, ratifica tratados, nombra altos funcionarios. Web: `senado.congreso.gob.pe`
- **Cámara de Diputados** (130 diputados) — cámara baja. Aprueba el presupuesto, inicia la mayoría de proyectos de ley, interpela ministros. Web: `diputados.congreso.gob.pe`
- **Pleno del Congreso** — sesión conjunta de ambas cámaras. Web: `bicameral.congreso.gob.pe`
- **SPLEY** sigue siendo el sistema de proyectos de ley. El periodo parlamentario aún es `2021` mientras el sistema se actualiza al nuevo ciclo.
- Un proyecto de ley ahora tiene que pasar por ambas cámaras para convertirse en ley (salvo excepciones del Reglamento).
- Las sesiones de todas las cámaras se publican en `comunicaciones.congreso.gob.pe/agenda`.

## Nomenclatura SPLEY

- Sufijos: `-CR` Congreso · `-PE` Poder Ejecutivo · `-GR` Gobierno Regional · `-GL` Gobierno Local y otros.
- Un `-PE` se mueve más rápido políticamente que un `-CR` de bancada chica.
- URL de expediente: `#/expediente/{PERIODO}/{NUMERO}` — el periodo es `2021`, no el año del número del proyecto.

## Formato

- Tablas para datos comparables. Prosa para análisis y conversación.
- Negritas solo para lo crítico.
- En expedientes: siempre cierra con **"Mi lectura"** de criterio propio.
- La extensión la decide el contenido, nunca el relleno. Una respuesta de una línea puede ser perfecta."""


# System prompt compacto para Phase 3 con tool results — ahorra ~800 tokens vs SYSTEM_BASE
SYSTEM_MINI = """Eres Lex, asistente parlamentario de Julio César en Perú. Fecha actual: {hoy}.
Reglas: datos exactos de la herramienta, nunca inventar URLs ni números. Tablas para datos comparables. Español directo, sin relleno. Negritas solo para lo crítico."""

# Bloques de formato inyectados en Fase 3 según la herramienta usada.
WORKFLOWS = {
    "fetch_expediente": """
⚠️⚠️ REGLAS ABSOLUTAS ANTES DE RESPONDER ⚠️⚠️

1. CADA DATO que escribas debe existir TEXTUALMENTE en el resultado de la herramienta. Si un campo no vino, escribe "—" o "Sin registros." NUNCA lo inventes, nunca lo deduzcas, nunca lo completes con tu conocimiento.
2. CADA URL/link debe ser copia EXACTA de lo que devolvió la herramienta. Si no hay URL, escribe `-`. JAMÁS construyas ni modifiques un link.
3. TODAS las filas de TODAS las tablas van siempre. Si la tabla tiene 20 filas, van las 20. Nada de "entre otros" ni resúmenes.
4. Las 5 pestañas van SIEMPRE en la respuesta, aunque estén vacías. Vacío es información.
5. Si el campo "seguimiento", "proyectos_acumulados", "documentacion_anexa", "secciones" o "opinion_ciudadana" vino vacío o ausente en los datos, escribe la sección con "Sin registros." — NO es lo mismo que inventar que hay datos.

---

## Expediente [**[campo: numero]**([campo: enlace_expediente]) — copiar exacto con sufijo, ej. 14864/2025-CR]
### [campo: titulo — en MAYÚSCULAS tal como viene]

---

### FICHA DEL PROYECTO

| CAMPO | DATO |
|---|---|
| Número | [campo: numero] |
| Título | [campo: titulo] |
| Sumilla | [campo: sumilla — si es igual al título, igual se pone] |
| Fecha de presentación | [campo: fecha_presentacion en dd/mm/aaaa] |
| Estado procesal | [campo: estado — EN MAYÚSCULAS] |
| Proponente | [campo: proponente] |
| Autores | [campo: autor_principal — lista COMPLETA, sin abreviar] |
| Coautores | [campo: coautores — si está vacío: —] |
| Adherentes | [campo: adherentes — si está vacío: —] |
| Grupo parlamentario | [campo: grupo_parlamentario — si está vacío: —] |
| Periodo parlamentario | [campo: periodo_parlamentario] |
| Legislatura | [campo: legislatura — si está vacío: —] |
| Link al expediente | [[Ver en SPLEY](campo: enlace_expediente — copiar EXACTO)] |

---

### AVANCE EN EL TRÁMITE

`Presentado → Enviado a Comisiones → En Comisiones → Debate en Pleno → Enviado al Ejecutivo → Ley Publicada`

**Etapa actual: [inferir de campo estado]** — [una línea explicando qué significa en la práctica: si está en comisiones, cuánto falta para el pleno; si está dormido, decirlo]

---

### COMISIONES A LAS QUE FUE DERIVADO

[Para CADA elemento del campo "comisiones":]
- **[comision.nombre]** — derivado el [comision.fecha_derivacion] [[Ver comisión](comision.enlace si existe)]

[Si el campo "comisiones" está vacío: "Todavía no ha sido derivado a ninguna comisión."]

---

### ACTOS DE TRABAJO POR COMISIÓN

⚠️ REGLA: Esta sección responde directamente "¿qué hizo cada comisión con este proyecto?". Para CADA comisión en el campo "comisiones", lista sus actos filtrando de "seguimiento" las filas donde seguimiento[i].comision coincide con el nombre de esa comisión. Orden cronológico ascendente (más antiguo primero).

[Para CADA comisión en el campo "comisiones":]

#### [comision.nombre]
*Derivado el [comision.fecha_derivacion]*

| FECHA | ACTO | DETALLE | ADJUNTOS |
|---|---|---|---|

Para cada fila de seguimiento donde seguimiento[i].comision == nombre de esta comisión:
- FECHA: seguimiento[i].fecha
- ACTO: seguimiento[i].estado — EN MAYÚSCULAS
- DETALLE: seguimiento[i].detalle — EN MAYÚSCULAS — si vacío: —
- ADJUNTOS: para CADA elemento de seguimiento[i].adjuntos → `[PDF](url)`. Si vacío: `-`

Si no hay ninguna fila de seguimiento que corresponda a esta comisión: escribir una sola línea → "Sin actos registrados en esta comisión aún."

[Si el campo "comisiones" está vacío: omitir esta sección completa — el seguimiento plano ya cubre todo.]

---

### 1. SEGUIMIENTO

⚠️ REGLA: Una fila por CADA elemento del campo "seguimiento". Si hay 1 elemento, 1 fila. Si hay 15, 15 filas. CERO excepciones. Orden: el más antiguo arriba (índice 0 primero).

| FECHA | ESTADO PROCESAL | COMISIÓN | DETALLE | ADJUNTOS |
|---|---|---|---|---|

Para cada fila:
- FECHA: campo seguimiento[i].fecha
- ESTADO PROCESAL: campo seguimiento[i].estado — EN MAYÚSCULAS
- COMISIÓN: campo seguimiento[i].comision — si está vacío: `—`
- DETALLE: campo seguimiento[i].detalle — EN MAYÚSCULAS — si está vacío: `—`
- ADJUNTOS: para CADA elemento de seguimiento[i].adjuntos → `[PDF](url)` o `[Oficio](url)` según tipo. Si hay varios, van todos separados por espacio. Si adjuntos está vacío o es lista vacía: `-`

[Si el campo "seguimiento" está vacío o ausente: escribir "Sin registros de seguimiento." y continuar con las demás secciones.]

---

### 2. PROYECTOS ACUMULADOS

⚠️ REGLA: Una fila por CADA elemento del campo "proyectos_acumulados". Todos van, sin excepción.

| N° PROYECTO | TÍTULO | FECHA PRESENTACIÓN | AUTOR | ESTADO | LINK |
|---|---|---|---|---|---|

Para cada fila:
- N° PROYECTO: campo proyectos_acumulados[i].numero
- TÍTULO: campo proyectos_acumulados[i].titulo
- FECHA PRESENTACIÓN: campo proyectos_acumulados[i].fecha_presentacion
- AUTOR: campo proyectos_acumulados[i].autor — si vacío: —
- ESTADO: campo proyectos_acumulados[i].estado
- LINK: `[Ver](`proyectos_acumulados[i].enlace`)` — si no hay enlace: `-`

[Si el campo "proyectos_acumulados" está vacío: "Sin proyectos acumulados registrados."]
[Si SÍ hay acumulados: añadir al final → **Ojo:** este proyecto tiene [N] acumulados, lo que indica respaldo de varias bancadas y mayor probabilidad de avance.]

---

### 3. DOCUMENTACIÓN ANEXA

⚠️ REGLA: Una fila por CADA elemento del campo "documentacion_anexa". Todos van.

| FECHA | TIPO DE DOCUMENTO | DESCRIPCIÓN / REMITENTE | ADJUNTOS |
|---|---|---|---|

Para cada fila:
- FECHA: campo documentacion_anexa[i].fecha — si vacío: —
- TIPO: campo documentacion_anexa[i].tipo — si vacío: —
- DESCRIPCIÓN: campo documentacion_anexa[i].descripcion — si vacío: —
- ADJUNTOS: para CADA archivo en documentacion_anexa[i].adjuntos → `[PDF](url)`. Si la lista está vacía: `-`

[Si "documentacion_anexa" está vacío o ausente: "Sin documentación anexa registrada."]
[Si la herramienta no devolvió esta pestaña (campo ausente del todo): "La herramienta no devolvió datos de Documentación Anexa — puede ser un problema de scraping. Verificar manualmente en el link del expediente."]

---

### 4. SECCIONES

⚠️ REGLA: Una fila por CADA elemento del campo "secciones". Todos van.

| SECCIÓN | CONTENIDO / RESUMEN | ADJUNTOS |
|---|---|---|

Para cada fila:
- SECCIÓN: campo secciones[i].titulo
- CONTENIDO: si secciones[i].texto tiene contenido → primeras 300 caracteres del texto + "…" si hay más. Si está vacío: —
- ADJUNTOS: para CADA archivo en secciones[i].adjuntos → `[PDF](url)`. Si la lista está vacía o ausente: `-`

[Si "secciones" está vacío: "Sin secciones registradas."]

---

### 5. OPINIÓN CIUDADANA

[Si opinion_ciudadana tiene datos con total_opiniones > 0:]

**Total de opiniones registradas: [opinion_ciudadana.total_opiniones]**
| A FAVOR | EN CONTRA | COMENTARIOS |
|---|---|---|
| [opinion_ciudadana.a_favor] | [opinion_ciudadana.en_contra] | [opinion_ciudadana.comentarios] |

[Si total_opiniones es 0 o el campo está vacío: "Sin opiniones ciudadanas registradas."]

---

### PREDICTAMEN / DICTAMEN

[Si el campo "predictamen" no es null:]
**Hay predictamen:** [predictamen.nombre] — fecha: [predictamen.fecha] — [[Ver documento](predictamen.url)]
[Aclarar si es favorable, desfavorable o con texto sustitutorio si el nombre lo indica]

[Si predictamen es null: "No hay predictamen ni dictamen registrado en ninguna comisión."]

---

### MI LECTURA

[Análisis propio de 4-8 líneas con criterio real — esto SÍ puede venir de tu razonamiento:]
- Estado real del proyecto: ¿avanzando o dormido? Si la fecha del último movimiento en seguimiento fue hace más de 60 días, está dormido — dilo explícitamente con la fecha.
- ¿Qué le falta para llegar al Pleno? (dictamen, debate en comisión, etc.)
- ¿Quién lo impulsa (bancada/proponente) y qué peso político tiene eso? (un -PE del Ejecutivo tiene más tracción que un -CR de bancada chica)
- ¿A qué sector afecta? ¿Por qué le importa al usuario?
- Si hay acumulados: ¿el respaldo de múltiples bancadas cambia el panorama?

---

### ADJUNTOS DEL EXPEDIENTE

⚠️ REGLA: Lista TODOS los adjuntos del campo "todos_los_adjuntos". Cada adjunto en una línea con link clickeable. Si el campo está vacío o ausente: "Sin adjuntos registrados."

Para cada elemento en todos_los_adjuntos:
- 📄 **[todos_los_adjuntos[i].nombre]** — [todos_los_adjuntos[i].descripcion si es distinta al nombre, si no omitir] → [Descargar PDF](todos_los_adjuntos[i].url)

[Si "todos_los_adjuntos" está vacío: "Sin adjuntos registrados en el expediente."]

---

**[Ver expediente completo en SPLEY]([campo: enlace_expediente])**""",

    "fetch_agenda_comisiones": """
## Formato para AGENDA DE COMISIONES

⚠️ REGLA ABSOLUTA: Muestra SOLO las sesiones que están literalmente en los datos devueltos. Si la herramienta devuelve sin_datos, vacío o error → escribe: "No hay sesiones de comisiones programadas para los próximos 2 días." JAMÁS inventes fechas, nombres de comisiones, ni números de proyectos.

Si hay datos reales:
| Fecha | Hora | Comisión | Lugar / Modalidad | Link a la agenda |
|---|---|---|---|---|
| dd/mm | HH:MM | [Comisión] | [Sala X / Virtual] | [URL exacta devuelta por la herramienta] |

Ordena por fecha y hora. Si un campo no vino, pon "—". Solo incluyes links que la herramienta devolvió — jamás construyas una URL. Cierra con criterio: qué sesión conviene seguir.""",

    "fetch_agenda_pleno": """
## Formato para AGENDA DEL PLENO

⚠️ REGLA ABSOLUTA: Usa SOLO los datos del documento real devuelto. Si no hay agenda o la herramienta devuelve vacío/error → escribe: "No hay Agenda del Pleno publicada para esta semana." JAMÁS inventes dictámenes, mociones ni números de proyecto.

Si hay datos reales:
## Agenda del Pleno — [fecha de la agenda]

### Resumen en números
| Tipo | Cantidad |
|---|---|
| Dictámenes | X |
| Denuncias constitucionales | X |
| Mociones | X |
| Insistencias / observadas | X |
| [Otros tipos que aparezcan] | X |

### Lo más relevante
- [3-5 puntos concretos de la agenda con número de proyecto/dictamen]

### Mi lectura
[Qué tiene pinta de votarse primero, qué es lo políticamente caliente]""",

    "fetch_interpelaciones": """
## Formato para INTERPELACIONES
## Interpelaciones — [fecha de hoy]

### Mociones presentadas formalmente
| Ministro | Cartera | Fecha de presentación | Estado | Motivo (resumen) |
|---|---|---|---|---|

Si no hay ninguna: "No hay mociones de interpelación presentadas formalmente ahorita."

### En gestación (según prensa)
[Para cada noticia en noticias_prensa que mencione interpelación o moción:]
- [Ministro X]: [medio] reporta que [bancada] está juntando firmas por [motivo]. Fuente: [URL exacta de la herramienta].
Si no hay: "Tampoco encontré noticias de firmas en curso."

### Mi lectura
[¿Alguna tiene los votos (se necesitan 25 firmas)? ¿Es presión política real o maniobra mediática? Análisis directo en 2-3 líneas.]

Distingue SIEMPRE lo formal (dato del sistema del Congreso) de lo periodístico (prensa). Nunca presentes un rumor de prensa como moción presentada.""",

    "fetch_agenda_camaras": """
## Formato para AGENDA BICAMERAL (Senado / Diputados / Pleno)

⚠️ REGLA ABSOLUTA: Muestra SOLO las sesiones que están en los datos devueltos. JAMÁS inventes sesiones, horas ni lugares. Si sin_datos=true → explica el posible motivo (feriado, receso, fin de semana) y sugiere alternativas concretas (agenda semanal, proyectos en trámite, destacados y citaciones).

Si hay datos:
## Agenda del Congreso — [fechas cubiertas]

| Fecha | Hora | Cámara | Sesión / Tema | Lugar |
|---|---|---|---|---|
| dd/mm | HH:MM | Senado / Diputados / Pleno | [tema] | [lugar] |

Ordena por fecha y hora. Agrupa visualmente por cámara si hay varias del mismo día.

### Mi lectura
[Qué sesión conviene seguir y por qué — en 2-3 líneas.]""",

    "buscar_proyectos": """
## Formato para PROYECTOS

⚠️ El número del proyecto SIEMPRE debe ser un link markdown usando el campo enlace: `[[numero](enlace)]`. Si no hay enlace: el número sin link.

| N° PROYECTO | FECHA | ESTADO | PROPONENTE | COMISIÓN | SUMILLA |
|---|---|---|---|---|---|
| [[numero](enlace)] | fecha | estado | proponente | comision | sumilla |

Máximo 15 filas. Si buscaste por materia y los resultados no corresponden al tema, dilo.

### Detalle
Para cada proyecto en la tabla, agrega debajo una línea con formato limpio:
**[numero]** — [sumilla completa sin abreviar]. Autores: [autor si viene en los datos].

### Adjuntos
Si algún proyecto tiene campo enlace con PDF directo o archivos adjuntos, listarlos así:
- **[numero]:** [Texto del proyecto](url_pdf) | [Exposición de motivos](url) — según lo que devuelva la herramienta.
Si no hay adjuntos en los datos devueltos: omitir esta sección.

Si el usuario quiere el expediente completo de uno, sugerir usar fetch_expediente.""",
}

# Flujos que dependen de PDF/transcript cargado (no de una herramienta).
WORKFLOW_PDF_FORMULA = """
## Formato para FÓRMULA LEGAL DESDE PDF
Localiza la sección "Fórmula Legal" (o "Texto del Proyecto de Ley") del PDF. La exposición de motivos es solo contexto.

## PL [número] — [título]

### Qué propone (en cristiano)
[2-4 líneas sin jerga]

### Artículo por artículo
- **Art. 1:** [qué establece]
- **Disposiciones complementarias/finales/derogatorias:** [ojo, suelen esconder lo importante]

### 🔴 Modifica leyes vigentes
- Modifica el Art. X de la Ley N° XXXX: [qué cambia, antes vs. después]
[Si no: "No modifica ninguna ley vigente, es norma nueva."]

### 🟡 Cambia plazos o procedimientos
- [Plazo/procedimiento actual → propuesto]
[Si no: "No toca plazos ni procedimientos existentes."]

### Mi lectura
[A quién afecta, qué tan viable, qué sector debe estar atento]

Las secciones 🔴 y 🟡 son OBLIGATORIAS siempre, aunque sea para decir que no aplican."""

WORKFLOW_SESION = """
## Formato para ANÁLISIS DE SESIÓN (solo sobre el transcript disponible)
## Sesión: [comisión o Pleno] — [fecha si consta]

### Temas debatidos
- [Tema]: [quién lo sustentó, 1-2 líneas]

### Lo que se votó
| Tema / Proyecto | Resultado | Detalle |
|---|---|---|
| PL 1234 — [título] | Aprobado/Rechazado/Cuarto intermedio | [votos si constan, o "por unanimidad", o "no se detalló el conteo"] |

### Acuerdos sin votación
- [Consensos, pases a comisión, pedidos aceptados]

### Lo que quedó pendiente
- [Temas anunciados no tratados, votaciones postergadas]

Si el transcript no menciona votaciones: "En esta sesión se debatió pero no se votó nada." Nunca deduzcas un resultado no dicho textualmente."""


TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "buscar_proyectos",
            "description": (
                "Obtiene proyectos de ley del Congreso del Perú desde el sistema SPLEY. "
                "Úsala cuando el usuario pida proyectos, leyes, expedientes o quiera buscar "
                "por tema/materia, autor/congresista, comisión, número de proyecto o rango de fechas. "
                "Para búsquedas por TEMA usa el parámetro 'materia'. "
                "Para los ÚLTIMOS N DÍAS usa el parámetro 'dias' (ej: dias=15 para últimos 15 días). "
                "Para un número específico usa 'numero'. Para un autor usa 'autor'."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "materia": {
                        "type": "string",
                        "description": "Tema o materia a buscar (ej: 'educacion', 'salud', 'transporte', 'mineria')"
                    },
                    "autor": {
                        "type": "string",
                        "description": "Apellido o nombre del congresista autor del proyecto"
                    },
                    "comision": {
                        "type": "string",
                        "description": "Nombre de la comisión parlamentaria"
                    },
                    "numero": {
                        "type": "string",
                        "description": "Número del proyecto de ley (ej: '14860/2025-CR' o solo '14860')"
                    },
                    "legislatura": {
                        "type": "string",
                        "description": "Período legislativo (default: '2021-2026')"
                    },
                    "dias": {
                        "type": "integer",
                        "description": "Filtrar proyectos presentados en los últimos N días calendario (ej: 15 para los últimos 15 días)"
                    },
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_sesiones",
            "description": (
                "Obtiene sesiones del Congreso del Perú desde el visor oficial. "
                "Usa esta herramienta cuando el usuario pregunte por sesiones, debates, "
                "votaciones o reuniones de comisiones."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "comision": {
                        "type": "string",
                        "description": "Nombre de la comisión"
                    },
                    "fecha": {
                        "type": "string",
                        "description": "Fecha en formato YYYY-MM-DD"
                    },
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_agenda",
            "description": (
                "Obtiene la agenda parlamentaria actual del Congreso del Perú: "
                "convocatorias, fechas y horarios de próximas sesiones."
            ),
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_destacados",
            "description": (
                "Obtiene las noticias y citaciones destacadas del Congreso del Perú."
            ),
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_congresista",
            "description": (
                "Obtiene el perfil completo de un congresista: todos sus proyectos de ley "
                "presentados, resumen por estado (aprobado, en comisión, archivado, etc.) "
                "y noticias recientes sobre esa persona. "
                "Úsala cuando el usuario pregunte por un congresista específico, "
                "quiera saber qué ha legislado alguien, o necesite el historial de un parlamentario."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "nombre": {
                        "type": "string",
                        "description": "Nombre o apellido del congresista (ej: 'Montoya', 'Patricia Chirinos')"
                    }
                },
                "required": ["nombre"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "buscar_en_web",
            "description": (
                "Busca cualquier tema general en internet usando DuckDuckGo. "
                "Úsala para preguntas que NO son sobre proyectos de ley, sesiones, agenda o congresistas específicos: "
                "historia, definiciones, noticias generales, conceptos legales, datos del mundo, etc."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Términos de búsqueda"
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Número de resultados (default: 5)"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "rastrear_proyecto",
            "description": (
                "Obtiene el estado detallado y actual de un proyecto de ley específico "
                "por su número. Úsala cuando el usuario quiera saber en qué estado está "
                "un proyecto puntual, si fue aprobado, archivado, o está en comisión."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "numero": {
                        "type": "string",
                        "description": "Número del proyecto (ej: '1234/2024-CR' o simplemente '1234')"
                    }
                },
                "required": ["numero"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_expediente",
            "description": (
                "Obtiene el expediente COMPLETO de un proyecto de ley desde el portal SPLEY del "
                "Congreso, con sus 5 pestañas: (1) Seguimiento — todos los movimientos con fecha, "
                "estado procesal, comisión, detalle y adjuntos; (2) Proyectos Acumulados; "
                "(3) Documentación Anexa — oficios, opiniones de ministerios, informes; "
                "(4) Secciones — texto del proyecto, fórmula legal, dictámenes, autógrafas; "
                "(5) Opinión Ciudadana. Usar cuando el usuario pida el expediente, el seguimiento, "
                "el trámite en comisiones, los actos de trabajo, los adjuntos o el predictamen de "
                "un proyecto específico. Si el usuario solo dio el tema, primero identificar el "
                "número con buscar_proyectos."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "numero_proyecto": {
                        "type": "string",
                        "description": "Número del proyecto de ley. Acepta formato oficial completo ('14864/2025-CR') o solo el correlativo ('14864')."
                    },
                    "periodo": {
                        "type": "string",
                        "description": "Periodo parlamentario, ej. '2021' para el periodo 2021-2026. Por defecto usa el periodo vigente."
                    }
                },
                "required": ["numero_proyecto"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_agenda_comisiones",
            "description": (
                "Obtiene las sesiones de comisiones programadas para los próximos días desde "
                "la web del Congreso. Devuelve por cada sesión: fecha, hora, comisión, lugar o "
                "modalidad, y link a la agenda. Usar cuando el usuario pregunte qué sesiones de "
                "comisiones hay hoy, mañana o en los próximos días."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "dias": {
                        "type": "integer",
                        "description": "Cantidad de días hacia adelante a consultar. Por defecto 2."
                    },
                    "comision": {
                        "type": "string",
                        "description": "Opcional. Filtrar por nombre (o parte del nombre) de una comisión específica, ej. 'Energía y Minas'."
                    }
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_agenda_pleno",
            "description": (
                "Obtiene la estructura de la Agenda del Pleno vigente desde la web del Congreso: "
                "cuántos dictámenes, denuncias constitucionales, mociones e insistencias hay "
                "agendados, con el detalle de cada ítem. Usar cuando el usuario pregunte por la "
                "Agenda del Pleno, qué se va a debatir en el Pleno, o cuántos dictámenes/denuncias hay."
            ),
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "responder_directo",
            "description": (
                "Usar cuando la pregunta NO requiere datos actualizados del Congreso ni "
                "búsqueda web: saludos ('hola', 'buenos días'), preguntas sobre ti mismo, "
                "seguimiento de una respuesta anterior ('¿y eso qué implica?', 'explícame eso'), "
                "conceptos, definiciones o historia que puedes responder con tu conocimiento."
            ),
            "parameters": {
                "type": "object",
                "properties": {}
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_agenda_camaras",
            "description": (
                "Obtiene las sesiones programadas del Senado, Cámara de Diputados y Pleno del "
                "Congreso bicameral desde comunicaciones.congreso.gob.pe/agenda. "
                "Usar cuando el usuario pregunte por sesiones del Senado, Diputados, Pleno "
                "bicameral, o la agenda general del nuevo Congreso."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "dias": {
                        "type": "integer",
                        "description": "Días hacia adelante a consultar (default: 2)."
                    },
                    "camara": {
                        "type": "string",
                        "description": "Filtrar por cámara: 'senado', 'diputados', 'pleno', 'comision'. Omitir para ver todas."
                    }
                }
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_interpelaciones",
            "description": (
                "Obtiene las mociones de interpelación a ministros presentadas formalmente ante "
                "el Congreso y noticias sobre mociones en gestación (recolección de firmas). "
                "Devuelve por cada moción: ministro, cartera, fecha, estado y motivo. "
                "Usar cuando el usuario pregunte por interpelaciones o mociones contra ministros. "
                "IMPORTANTE: complementar siempre con buscar_en_web para detectar mociones en "
                "recolección de firmas que aún no aparecen en el sistema."
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "ministro": {
                        "type": "string",
                        "description": "Opcional. Filtrar por nombre del ministro o de la cartera, ej. 'Interior' o 'Ministro de Salud'."
                    }
                }
            }
        }
    }
]

async def buscar_en_web(query: str, limit: int = 5):
    try:
        loop = __import__('asyncio').get_event_loop()
        def _search():
            with DDGS() as ddgs:
                results = list(ddgs.text(query, max_results=limit))
            return results
        results = await loop.run_in_executor(None, _search)
        return [{"titulo": r.get("title"), "url": r.get("href"), "resumen": r.get("body")} for r in results]
    except Exception as e:
        return {"sin_datos": True, "mensaje": str(e)}

TOOL_MAP = {
    "buscar_proyectos":        lambda args: fetch_proyectos(**args),
    "buscar_sesiones":         lambda args: fetch_sesiones(**args),
    "buscar_agenda":           lambda args: fetch_agenda(),
    "buscar_destacados":       lambda args: fetch_destacados(),
    "buscar_congresista":      lambda args: fetch_congresista(**args),
    "rastrear_proyecto":       lambda args: fetch_estado_proyecto(**args),
    "buscar_en_web":           lambda args: buscar_en_web(**args),
    "fetch_expediente":        lambda args: fetch_expediente(
                                   numero=args.get("numero_proyecto") or args.get("numero", "")
                               ),
    "fetch_agenda_comisiones": lambda args: fetch_agenda_comisiones(**{k: v for k, v in args.items() if k in ("dias", "comision")}),
    "fetch_agenda_pleno":      lambda args: fetch_agenda_pleno(),
    "fetch_agenda_camaras":    lambda args: fetch_agenda_camaras(**{k: v for k, v in args.items() if k in ("dias", "camara")}),
    "fetch_interpelaciones":   lambda args: fetch_interpelaciones(**{k: v for k, v in args.items() if k in ("ministro",)}),
    "responder_directo":       lambda args: _responder_directo(),
}

async def _responder_directo():
    return {"nota": "Responde directamente con tu conocimiento, sin datos externos."}

STATUS_LABELS = {
    "buscar_proyectos":        "Buscando proyectos de ley en SPLEY...",
    "buscar_sesiones":         "Consultando sesiones del Congreso...",
    "buscar_agenda":           "Obteniendo agenda parlamentaria...",
    "buscar_destacados":       "Cargando noticias del Congreso...",
    "buscar_congresista":      "Consultando perfil del congresista...",
    "rastrear_proyecto":       "Rastreando estado del proyecto...",
    "buscar_en_web":           "Buscando en internet...",
    "fetch_expediente":        "Consultando el expediente completo en SPLEY (5 pestañas)...",
    "fetch_agenda_comisiones": "Revisando agenda de comisiones...",
    "fetch_agenda_pleno":      "Cargando la Agenda del Pleno...",
    "fetch_agenda_camaras":    "Revisando agenda del Congreso bicameral...",
    "fetch_interpelaciones":   "Buscando mociones de interpelación...",
    "responder_directo":       "Pensando...",
}


@app.get("/status")
async def status():
    ready = bool(os.getenv("GROQ_API_KEY", ""))
    return {"ready": ready}


@app.get("/", response_class=HTMLResponse)
async def root():
    return (Path("static") / "index.html").read_text()


@app.get("/static/sw.js")
async def service_worker():
    return FileResponse("static/sw.js", media_type="application/javascript",
                        headers={"Service-Worker-Allowed": "/"})


@app.get("/sessions", response_class=HTMLResponse)
async def sessions_page():
    return (Path("static") / "sessions.html").read_text()


@app.get("/pdfs", response_class=HTMLResponse)
async def pdfs_page():
    return (Path("static") / "pdfs.html").read_text()


REFERENCIAS_PDF = [
    {
        "titulo": "Reglamento del Congreso de la República (setiembre 2025)",
        "enlace": "https://www3.congreso.gob.pe/Docs/constitucion/reglamento/reglamento%20setiembre-2025.pdf",
        "tipo": "Referencia",
    },
    {
        "titulo": "Constitución Política del Perú (dic. 2024)",
        "enlace": "https://www3.congreso.gob.pe/Docs/files/constitucion/constitucion-12-2024.pdf",
        "tipo": "Referencia",
    },
    {
        "titulo": "Manual de Técnica Legislativa — 3ra edición",
        "enlace": "https://www3.congreso.gob.pe/Docs/dgp/files/manual-tecnica-legislativa-3raedicion.pdf",
        "tipo": "Referencia",
    },
]

@app.get("/congreso-pdfs")
async def congreso_pdfs():
    """PDFs rápidos: destacados del homepage + referencias fijas."""
    pdfs = []
    try:
        data = await fetch_destacados()
        for item in data.get("destacados", []):
            url = item.get("enlace", "")
            if url.lower().endswith(".pdf"):
                pdfs.append({"titulo": item["titulo"], "enlace": url, "tipo": "Destacado"})
        for item in data.get("citaciones", []):
            url = item.get("enlace", "")
            if url.lower().endswith(".pdf"):
                pdfs.append({"titulo": item["titulo"], "enlace": url, "tipo": "Citación"})
    except Exception:
        pass
    seen = {p["enlace"] for p in pdfs}
    for ref in REFERENCIAS_PDF:
        if ref["enlace"] not in seen:
            seen.add(ref["enlace"])
            pdfs.append(ref)
    return {"pdfs": pdfs}

@app.get("/congreso-proyectos")
async def congreso_proyectos():
    """Proyectos SPLEY recientes — se carga en segundo plano."""
    try:
        data = await fetch_proyectos(limit=15)
        proyectos = []
        for item in data.get("items", []):
            numero = item.get("numero", "")
            titulo = item.get("sumilla", numero)
            enlace = item.get("enlace", "")
            if enlace:
                proyectos.append({
                    "titulo": f"[{numero}] {titulo[:100]}" if numero else titulo[:110],
                    "enlace": enlace,
                    "tipo": "Proyecto de Ley",
                })
        return {"pdfs": proyectos}
    except Exception:
        return {"pdfs": []}


@app.get("/pdf-thumbnail")
async def pdf_thumbnail(url: str = Query(...)):
    import fitz
    try:
        async with httpx.AsyncClient(timeout=20, follow_redirects=True) as c:
            r = await c.get(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/124.0.0.0"})
        if r.status_code != 200 or "pdf" not in r.headers.get("content-type", "application/pdf").lower():
            ct = r.headers.get("content-type", "")
            if "html" in ct or r.status_code != 200:
                return Response(status_code=404, content=b"not a pdf")
        doc = fitz.open(stream=r.content, filetype="pdf")
        page = doc[0]
        mat = fitz.Matrix(1.8, 1.8)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        png_bytes = pix.tobytes("png")
        return Response(content=png_bytes, media_type="image/png")
    except Exception as e:
        return Response(status_code=400, content=str(e).encode())


@app.post("/load-pdf-url")
async def load_pdf_url(request: Request):
    import fitz
    body = await request.json()
    url  = body.get("url", "")
    try:
        async with httpx.AsyncClient(timeout=30, follow_redirects=True) as c:
            r = await c.get(url, headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) Chrome/124.0.0.0"})
        doc  = fitz.open(stream=r.content, filetype="pdf")
        pages = len(doc)
        text  = "\n\n".join(page.get_text() for page in doc).strip()
        if len(text) > 40000:
            text = text[:40000] + f"\n\n[Texto recortado — documento original: {pages} páginas]"
        name = url.split("/")[-1]
        return {"ok": True, "pages": pages, "text": text, "filename": name}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.post("/upload-pdf")
async def upload_pdf(file: UploadFile = File(...)):
    import fitz
    try:
        data = await file.read()
        doc  = fitz.open(stream=data, filetype="pdf")
        pages = len(doc)
        text  = "\n\n".join(page.get_text() for page in doc).strip()
        # Cap en 40000 chars para no reventar el contexto
        if len(text) > 40000:
            text = text[:40000] + f"\n\n[Texto recortado — documento original: {pages} páginas]"
        return {"ok": True, "pages": pages, "text": text, "filename": file.filename}
    except Exception as e:
        return {"ok": False, "error": str(e)}


@app.post("/chat")
async def chat(request: Request):
    body     = await request.json()
    messages = body.get("messages", [])
    api_key  = GROQ_API_KEY

    if not api_key:
        async def err():
            yield f"data: {json.dumps({'error': 'Falta la API key de Groq'})}\n\n"
        return StreamingResponse(err(), media_type="text/event-stream")

    client = Groq(api_key=api_key)

    async def generate():
        # Inyectar fecha actual en el system prompt para evitar alucinaciones temporales
        hoy = datetime.now().strftime("%d/%m/%Y")
        system_con_fecha = SYSTEM_BASE + f"\n\n**Fecha actual: {hoy}** — Solo muestra sesiones o eventos a partir de hoy. Si una herramienta no devuelve sesiones reales, NO digas simplemente 'no hay sesiones'. En cambio: (1) explica brevemente el posible motivo (feriado, receso parlamentario, fin de semana, etc. según la fecha), (2) sugiere alternativas concretas como revisar la agenda de la semana siguiente, consultar proyectos de ley en trámite, o revisar los destacados y citaciones. Sé directo y útil, no te limites a dar una negativa seca."

        def _parse_retry_seconds(e) -> float:
            """Extrae los segundos de espera del error de rate limit de Groq."""
            s = str(e)
            # "Please try again in 2.5s" o "try again in 750ms"
            m = re.search(r"try again in ([0-9.]+)s", s, re.IGNORECASE)
            if m:
                return float(m.group(1)) + 0.5
            m = re.search(r"try again in ([0-9.]+)ms", s, re.IGNORECASE)
            if m:
                return float(m.group(1)) / 1000 + 0.5
            m = re.search(r"try again in (\d+)m(\d+(?:\.\d+)?)s", s, re.IGNORECASE)
            if m:
                return int(m.group(1)) * 60 + float(m.group(2)) + 0.5
            return 12.0  # fallback conservador

        def _friendly_error(e):
            s = str(e).lower()
            if "per day" in s or "tpd" in s:
                m = re.search(r"try again in ([0-9hms.]+)", s)
                cuando = "en un rato"
                if m:
                    mins = re.search(r"(\d+)m", m.group(1))
                    cuando = f"en ~{mins.group(1)} min" if mins else f"en {m.group(1)}"
                return (f"Llegamos al límite de tokens por ahora. Vuelve a intentar {cuando}.")
            if "rate limit" in s or "429" in s or "tokens per" in s or "quota" in s:
                return "Muchas consultas muy rápido. Espera unos segundos y vuelve a intentarlo."
            return "Hubo un problema al conectar. Intentá de nuevo."

        # Detectar si es solicitud de resumen semanal
        last_msg = messages[-1].get("content", "") if messages else ""
        is_resumen = last_msg.strip().startswith("__RESUMEN_SEMANAL__")
        sector = None
        if is_resumen and ":" in last_msg:
            sector = last_msg.strip().split(":", 1)[1].strip()

        # ¿Hay un PDF/documento cargado en el historial reciente? El frontend lo
        # inyecta como un mensaje que empieza con "He cargado el documento".
        recientes = messages[-6:]
        doc_en_contexto = any(
            "He cargado el documento" in (m.get("content", "") or "")
            for m in recientes if m.get("role") == "user"
        )
        low_last = last_msg.lower()
        pide_analisis = any(w in low_last for w in (
            "analiza", "analizar", "resume", "resumir", "resumen", "fórmula legal",
            "formula legal", "qué propone", "que propone", "qué modifica", "que modifica",
            "artículo", "articulo", "deroga", "explica",
        ))
        # Análisis de un documento ya cargado: se trabaja sobre el texto en
        # contexto, NO se necesita ninguna herramienta de scraping.
        analizar_documento = doc_en_contexto and (pide_analisis or len(last_msg) < 60)

        # Detectar link de sesión / transcript para el flujo de análisis de sesión.
        has_sesion = ("youtube.com" in low_last or "youtu.be" in low_last
                      or "transcript" in low_last or "[sesión" in low_last)
        has_pdf    = analizar_documento

        # conversation: historial para la Fase 3. El system se arma dinámicamente
        # (base compacta + solo los flujos relevantes) para no reventar tokens.
        if is_resumen:
            base_msg = "Genera el resumen ejecutivo semanal completo del Congreso del Perú."
            if sector and sector != "general":
                base_msg += f" Enfoca el análisis especialmente en el sector {sector} y los proyectos de ley, noticias y agenda que impacten a ese sector."
            conversation = [{"role": "user", "content": base_msg}]
        else:
            # Mantener historial amplio para conversación fluida
            conversation = messages[-20:]

        # Short-circuit: analizar un documento cargado o una sesión no requiere
        # scraping. Vamos directo a la Fase 3 con el flujo correspondiente.
        if analizar_documento or (has_sesion and not is_resumen):
            system_p3 = system_con_fecha
            if analizar_documento:
                system_p3 += "\n" + WORKFLOW_PDF_FORMULA
            if has_sesion:
                system_p3 += "\n" + WORKFLOW_SESION
                # Intentar obtener el transcript del video de YouTube antes de Fase 3
                yt_match = re.search(
                    r"(?:youtube\.com/(?:watch\?v=|live/|embed/)|youtu\.be/)([A-Za-z0-9_-]{11})",
                    last_msg
                )
                if yt_match:
                    video_id = yt_match.group(1)
                    yield f"data: {json.dumps({'status': 'Obteniendo transcript de la sesión...'})}\n\n"
                    try:
                        transcript_data = await fetch_transcript_youtube(video_id)
                    except Exception:
                        transcript_data = None

                    if transcript_data and transcript_data.get("ok") and transcript_data.get("text"):
                        # Insertar el transcript como contexto justo antes del último mensaje del usuario
                        transcript_msg = {
                            "role": "user",
                            "content": (
                                f"[TRANSCRIPT DE LA SESIÓN — fuente: {transcript_data.get('source', 'youtube')}]\n"
                                f"{transcript_data['text']}\n"
                                "[FIN DEL TRANSCRIPT]"
                            )
                        }
                        conversation = conversation[:-1] + [transcript_msg, conversation[-1]]
                    else:
                        # Sin subtítulos disponibles — informar directamente sin pasar por el modelo
                        yield f"data: {json.dumps({'text': 'No pude obtener el transcript de ese video (sin subtítulos disponibles o video privado). Para analizar la sesión podés: (1) cargar el PDF del acta con el botón de adjunto, o (2) pegar el texto del transcript directamente en el chat.'})}\n\n"
                        yield "data: [DONE]\n\n"
                        return

            msgs_directo = [{"role": "system", "content": system_p3}] + conversation
            try:
                stream = client.chat.completions.create(
                    model="llama-3.3-70b-versatile",
                    messages=msgs_directo,
                    max_tokens=2048,
                    temperature=0.4,
                    stream=True,
                )
                for chunk in stream:
                    delta = chunk.choices[0].delta.content
                    if delta:
                        yield f"data: {json.dumps({'text': delta})}\n\n"
                yield "data: [DONE]\n\n"
            except Exception as e:
                yield f"data: {json.dumps({'error': _friendly_error(e)})}\n\n"
            return

        # router_msgs: prompt compacto solo para elegir tools en la Fase 1.
        # Si hay un doc gigante en contexto, no lo mandamos al router (gasta tokens
        # y no aporta a la elección de herramienta): usamos solo el texto del pedido.
        # Detectar si hay un expediente completo en el historial reciente
        recent_assistant = " ".join(
            m.get("content", "") for m in messages[-6:]
            if m.get("role") == "assistant"
        )
        has_expediente_en_contexto = any(
            marker in recent_assistant
            for marker in ("FICHA DEL PROYECTO", "COMISIONES A LAS QUE FUE DERIVADO",
                           "ACTOS DE TRABAJO POR COMISIÓN", "MI LECTURA")
        )

        if doc_en_contexto:
            router_msgs = [{"role": "system", "content": ROUTER_PROMPT},
                           {"role": "user", "content": last_msg}]
        elif has_expediente_en_contexto:
            # El expediente ya está en contexto: enviar solo la pregunta + nota breve.
            # NO mandar el historial completo al router — el expediente puede ser miles
            # de tokens y revienta el rate limit solo en Phase 1.
            router_context = (
                "[CONTEXTO: El asistente ya mostró el expediente completo de un proyecto "
                "en esta conversación (FICHA DEL PROYECTO, SEGUIMIENTO, COMISIONES, etc.). "
                "Si el usuario pregunta sobre ese expediente — comisiones, actos, predictamen, "
                "adjuntos, estado, autores — usa responder_directo. "
                "Solo llama fetch_expediente si pide OTRO proyecto diferente.]\n\n"
                f"Pregunta: {last_msg}"
            )
            router_msgs = [{"role": "system", "content": ROUTER_PROMPT},
                           {"role": "user", "content": router_context}]
        else:
            router_msgs = [{"role": "system", "content": ROUTER_PROMPT}] + messages[-4:]

        # ── Phase 1: let model decide if it needs tools ────────
        # Usar modelo pequeño (8B) para el router: solo elige una función,
        # no necesita 70B. El 8B tiene 20k TPM vs 6k del 70B — evita rate limits.
        ROUTER_MODEL = "llama-3.1-8b-instant"
        MAIN_MODEL   = "llama-3.3-70b-versatile"

        def _is_tool_format_error(e):
            s = str(e)
            return "tool_use_failed" in s or "failed_generation" in s or "400" in s

        try:
            resp = client.chat.completions.create(
                model=ROUTER_MODEL,
                messages=router_msgs,
                tools=TOOLS,
                tool_choice="required",
                max_tokens=512,
                temperature=0.2,
                stream=False,
            )
        except Exception as e:
            if _is_tool_format_error(e):
                # Model generated malformed tool call — retry without tools
                try:
                    resp2 = client.chat.completions.create(
                        model=MAIN_MODEL,
                        messages=[{"role": "system", "content": system_con_fecha}] + conversation,
                        max_tokens=2048,
                        temperature=0.4,
                        stream=True,
                    )
                    for chunk in resp2:
                        delta = chunk.choices[0].delta.content
                        if delta:
                            yield f"data: {json.dumps({'text': delta})}\n\n"
                    yield "data: [DONE]\n\n"
                except Exception as e2:
                    yield f"data: {json.dumps({'error': _friendly_error(e2)})}\n\n"
            else:
                yield f"data: {json.dumps({'error': _friendly_error(e)})}\n\n"
            return

        choice  = resp.choices[0]
        finish  = choice.finish_reason

        # ── Phase 2: execute tools if requested ────────────────
        tool_msgs   = []   # assistant tool_call + tool result messages
        tools_usados = []  # nombres de tools ejecutadas (para armar el flujo de Fase 3)
        solo_responder_directo = False  # señal para Phase 3 minimalista

        if finish == "tool_calls" and choice.message.tool_calls:
            # Ignorar responder_directo: es solo señal de "responde sin datos"
            real_calls = [tc for tc in choice.message.tool_calls
                          if tc.function.name != "responder_directo"]
            if not real_calls:
                solo_responder_directo = True

            if real_calls:
                tool_msgs.append({
                    "role": "assistant",
                    "content": choice.message.content or "",
                    "tool_calls": [
                        {
                            "id": tc.id,
                            "type": "function",
                            "function": {
                                "name": tc.function.name,
                                "arguments": tc.function.arguments,
                            }
                        }
                        for tc in real_calls
                    ]
                })

                for tc in real_calls:
                    name = tc.function.name
                    tools_usados.append(name)
                    try:
                        args = json.loads(tc.function.arguments or "{}")
                    except (ValueError, TypeError):
                        args = {}
                    if not isinstance(args, dict):
                        args = {}
                    if "limit" in args:
                        try:
                            args["limit"] = int(args["limit"])
                        except (ValueError, TypeError):
                            args["limit"] = 20
                    args = {k: v for k, v in args.items() if v != ""}

                    status = STATUS_LABELS.get(name, "Consultando el Congreso...")
                    yield f"data: {json.dumps({'status': status})}\n\n"

                    try:
                        if name not in TOOL_MAP:
                            result = {"sin_datos": True, "mensaje": f"Herramienta '{name}' no disponible."}
                        else:
                            result = await TOOL_MAP[name](args)
                            if isinstance(result, dict) and "error" in result:
                                result = {"sin_datos": True, "mensaje": result.get("error", "No hay información disponible.")}
                    except Exception as tool_err:
                        result = {"sin_datos": True, "mensaje": f"Error al consultar {name}: {str(tool_err)[:100]}"}

                    # Cap tool result: 7k chars ≈ ~2000 tokens — protege TPM del 8B
                    result_str = json.dumps(result, ensure_ascii=False)
                    if len(result_str) > 7000:
                        result_str = result_str[:7000] + '... [recortado]"}'
                    tool_msgs.append({
                        "role": "tool",
                        "tool_call_id": tc.id,
                        "content": result_str,
                    })

        # ── Build Phase 3 system prompt: base + solo los flujos relevantes ──
        if is_resumen:
            system_p3 = RESUMEN_PROMPT
        elif solo_responder_directo:
            system_p3 = system_con_fecha
        elif tool_msgs:
            # Con tool results: usar SYSTEM_MINI para ahorrar tokens (~800 menos)
            # fetch_expediente y PDF necesitan sus workflows completos igual
            mini = SYSTEM_MINI.format(hoy=hoy)
            has_heavy_workflow = any(t in WORKFLOWS for t in tools_usados) or has_pdf or has_sesion
            if has_heavy_workflow:
                system_p3 = system_con_fecha
                for t in tools_usados:
                    if t in WORKFLOWS:
                        system_p3 += "\n" + WORKFLOWS[t]
                if has_pdf:
                    system_p3 += "\n" + WORKFLOW_PDF_FORMULA
                if has_sesion:
                    system_p3 += "\n" + WORKFLOW_SESION
            else:
                system_p3 = mini
        else:
            system_p3 = system_con_fecha

        # Cuando hay tool results, recortar el historial enviado a Phase 3.
        # El tool result ya aporta el contexto; mandar 20 mensajes adicionales
        # dispara el TPM fácilmente. Con tools: solo los últimos 4 mensajes.
        conv_p3 = messages[-4:] if tool_msgs else conversation

        msgs = [{"role": "system", "content": system_p3}] + conv_p3 + tool_msgs

        # ── Phase 3: stream final answer ───────────────────────
        # Con tool results el request es pesado (>6k tokens) — usar 8B (20k TPM).
        # Sin tools (solo conversación) el request es ligero — usar 70B.
        _p3_model   = MAIN_MODEL  # 70B: 12k TPM vs 6k del 8B — siempre usar el de mayor límite
        if "fetch_expediente" in tools_usados:
            _max_tokens = 4000
        elif tools_usados:
            _max_tokens = 1800  # queries de proyectos/agenda: respuesta más corta, menos TPM
        else:
            _max_tokens = 2500

        import asyncio as _asyncio

        async def _stream_p3(msgs_in, max_tok):
            stream = client.chat.completions.create(
                model=_p3_model,
                messages=msgs_in,
                max_tokens=max_tok,
                temperature=0.4,
                stream=True,
            )
            for chunk in stream:
                delta = chunk.choices[0].delta.content
                if delta:
                    yield delta

        _is_rate = lambda e: any(x in str(e).lower() for x in ("rate limit","429","tokens per","quota","per day"))

        last_exc = None
        for attempt in range(3):
            try:
                async for delta in _stream_p3(msgs, _max_tokens):
                    yield f"data: {json.dumps({'text': delta})}\n\n"
                yield "data: [DONE]\n\n"
                last_exc = None
                break
            except Exception as e:
                last_exc = e
                if _is_rate(e) and attempt < 2:
                    wait = _parse_retry_seconds(e)
                    yield f"data: {json.dumps({'status': f'Límite de Groq, reintentando en {wait:.0f}s...'})}\n\n"
                    await _asyncio.sleep(wait)
                else:
                    break
        if last_exc:
            yield f"data: {json.dumps({'error': _friendly_error(last_exc)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/export/docx")
async def export_docx(request: Request):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    body = await request.json()
    md   = body.get("content", "")

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # Header line
    hdr_para = doc.add_paragraph()
    hdr_run  = hdr_para.add_run("DOCUMENTO CONFIDENCIAL — GESTIÓN DE ASUNTOS PÚBLICOS")
    hdr_run.font.size  = Pt(8)
    hdr_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    hdr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()  # spacer

    def strip_inline(text):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*',     r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return text

    def _add_hyperlink(paragraph, text, url):
        """Add a clickable hyperlink run to a paragraph."""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'Hyperlink')
        rPr.append(rStyle)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        hyperlink.append(r)
        paragraph._p.append(hyperlink)

    def add_md_para(para_text):
        """Add a paragraph with bold, italic and hyperlink support."""
        p = doc.add_paragraph()
        # Split on links first, then bold/italic
        tokens = re.split(r'(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|\*[^*]+\*)', para_text)
        for tok in tokens:
            link_m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
            if link_m:
                _add_hyperlink(p, link_m.group(1), link_m.group(2))
            elif tok.startswith('**') and tok.endswith('**'):
                run = p.add_run(tok[2:-2])
                run.bold = True
            elif tok.startswith('*') and tok.endswith('*'):
                run = p.add_run(tok[1:-1])
                run.italic = True
            else:
                p.add_run(tok)
        return p

    def is_table_row(s):
        return s.startswith('|') and s.endswith('|')

    def is_separator_row(s):
        return is_table_row(s) and re.match(r'^\|[\s\-|:]+\|$', s)

    def add_word_table(rows):
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        cols = len(rows[0])
        tbl  = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = 'Table Grid'
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                cell = tbl.cell(r_idx, c_idx)
                cell.text = strip_inline(cell_text.strip())
                if r_idx == 0:  # header row bold + dark bg
                    run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell.text)
                    cell.paragraphs[0].clear()
                    run = cell.paragraphs[0].add_run(cell_text.strip())
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd   = OxmlElement('w:shd')
                    shd.set(qn('w:val'),   'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'),  '1a1a1a')
                    tc_pr.append(shd)
        doc.add_paragraph()

    # Collect table rows before processing
    lines = md.split('\n')
    i = 0
    while i < len(lines):
        stripped = lines[i].rstrip()

        # Detect markdown table block
        if is_table_row(stripped) and i + 1 < len(lines) and is_separator_row(lines[i+1].rstrip()):
            parse_row = lambda s: [c for c in s.strip('|').split('|')]
            header = parse_row(stripped)
            i += 2  # skip header + separator
            data_rows = [header]
            while i < len(lines) and is_table_row(lines[i].rstrip()):
                data_rows.append(parse_row(lines[i].rstrip()))
                i += 1
            add_word_table(data_rows)
            continue

        if stripped.startswith('### '):
            doc.add_heading(strip_inline(stripped[4:]), level=3)
        elif stripped.startswith('## '):
            doc.add_heading(strip_inline(stripped[3:]), level=2)
        elif stripped.startswith('# '):
            doc.add_heading(strip_inline(stripped[2:]), level=1)
        elif stripped.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif re.match(r'^[-*]\s+', stripped):
            text   = re.sub(r'^[-*]\s+', '', stripped)
            p      = doc.add_paragraph(style='List Bullet')
            tokens = re.split(r'(\[[^\]]+\]\([^)]+\)|\*\*[^*]+\*\*|\*[^*]+\*)', text)
            for tok in tokens:
                link_m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', tok)
                if link_m:
                    _add_hyperlink(p, link_m.group(1), link_m.group(2))
                elif tok.startswith('**') and tok.endswith('**'):
                    run = p.add_run(tok[2:-2]); run.bold = True
                elif tok.startswith('*') and tok.endswith('*'):
                    run = p.add_run(tok[1:-1]); run.italic = True
                else:
                    p.add_run(tok)
        elif re.match(r'^\d+\.\s+', stripped):
            doc.add_paragraph(re.sub(r'^\d+\.\s+', '', stripped), style='List Number')
        elif stripped == '':
            doc.add_paragraph()
        else:
            add_md_para(stripped)
        i += 1

    # Footer
    doc.add_paragraph()
    ftr_para = doc.add_paragraph()
    date_str = datetime.now().strftime('%d/%m/%Y')
    ftr_run  = ftr_para.add_run(f"Generado por Lex — Sistema de Monitoreo Parlamentario · {date_str}")
    ftr_run.font.size  = Pt(8)
    ftr_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    ftr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"Resumen-Congreso-{datetime.now().strftime('%Y-%m-%d')}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.get("/sesiones/cookies-status")
async def sesiones_cookies_status():
    from scraper import _get_cookie_path, COOKIE_PATHS
    path = _get_cookie_path()
    return {"ok": bool(path), "path": path, "search_paths": COOKIE_PATHS}


@app.get("/sesiones/videos")
async def sesiones_videos():
    result = await fetch_videos_youtube(limit=15)
    return result


def _build_sesion_prompt(titulo: str, texto: str) -> str:
    return f"""Analiza este transcript de la sesión del Congreso del Perú titulada: "{titulo}".

Genera un resumen estructurado con este formato EXACTO:

## Resumen — {titulo}

### Temas tratados
[Tabla: Tema | Descripción | Resultado/Estado]

### Proyectos o normas mencionados
[Tabla: Número/Nombre | Tema | Posición mayoritaria]

### Votaciones o acuerdos
[Tabla: Asunto | A favor | En contra | Resultado]

### Puntos destacados
[Lista de los 3-5 momentos más relevantes de la sesión]

---
Transcript de la sesión:
{texto[:40000]}"""


@app.post("/sesiones/resumir")
async def sesiones_resumir(request: Request):
    body     = await request.json()
    video_id = body.get("video_id", "")
    titulo   = body.get("titulo", "este video")
    en_vivo  = body.get("en_vivo", False)
    api_key  = GROQ_API_KEY

    async def generate():
        if not video_id:
            yield f"data: {json.dumps({'error': 'Falta el ID del video'})}\n\n"
            return

        # ── Fase 1: subtítulos de YouTube ─────────────────────
        yield f"data: {json.dumps({'status': 'Buscando subtítulos en YouTube...'})}\n\n"
        loop = __import__('asyncio').get_event_loop()
        captions = await loop.run_in_executor(None, get_yt_captions, video_id)

        tr = captions  # can be None

        # ── Fase 2: Whisper si no hay subtítulos ──────────────
        if not tr:
            if not api_key:
                yield f"data: {json.dumps({'error': 'No hay subtítulos disponibles para este video.'})}\n\n"
                return

            minutes = 5 if en_vivo else 10
            label   = f"los últimos {minutes} min del stream en vivo" if en_vivo else f"los primeros {minutes} min"
            yield f"data: {json.dumps({'status': f'No hay subtítulos. Descargando audio ({label})... esto toma ~2 minutos.'})}\n\n"

            tr = await loop.run_in_executor(None, transcribe_with_whisper, video_id, api_key, minutes)
            if not tr.get("ok"):
                yield f"data: {json.dumps({'error': tr.get('error', 'No se pudo transcribir el audio.')})}\n\n"
                return

            nota = tr.get("nota", "")
            yield f"data: {json.dumps({'status': f'Audio transcrito. Analizando ({nota})...'})}\n\n"
        else:
            yield f"data: {json.dumps({'status': 'Subtítulos obtenidos. Analizando sesión...'})}\n\n"

        # Emitir transcript raw para que el frontend ofrezca descarga
        yt_url = f"https://www.youtube.com/watch?v={video_id}"
        yield f"data: {json.dumps({'transcript_raw': tr['text'], 'video_url': yt_url, 'video_titulo': titulo})}\n\n"

        prompt = _build_sesion_prompt(titulo, tr['text'])
        client = Groq(api_key=api_key)
        try:
            stream = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": "Eres Lex, experto en análisis parlamentario del Congreso del Perú. Analiza transcripts de sesiones y los conviertes en resúmenes ejecutivos con tablas."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=3000,
                temperature=0.3,
                stream=True,
            )
            for chunk in stream:
                delta = chunk.choices[0].delta.content
                if delta:
                    yield f"data: {json.dumps({'text': delta})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error("sesiones_resumir LLM error: %s", e)
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


@app.post("/sesiones/resumir-texto")
async def sesiones_resumir_texto(request: Request):
    """Resume un transcript que el usuario pegó manualmente."""
    body    = await request.json()
    texto   = body.get("texto", "").strip()
    titulo  = body.get("titulo", "esta sesión")
    api_key = GROQ_API_KEY

    async def generate():
        if not texto:
            yield f"data: {json.dumps({'error': 'No hay texto para resumir.'})}\n\n"
            return
        if not api_key:
            yield f"data: {json.dumps({'error': 'Falta la API key de Groq.'})}\n\n"
            return

        yield f"data: {json.dumps({'status': 'Analizando transcripción...'})}\n\n"

        prompt = _build_sesion_prompt(titulo, texto)
        client = Groq(api_key=api_key)
        try:
            stream = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": "Eres Lex, experto en análisis parlamentario del Congreso del Perú."},
                    {"role": "user", "content": prompt},
                ],
                max_tokens=3000,
                temperature=0.3,
                stream=True,
            )
            for chunk in stream:
                delta = chunk.choices[0].delta.content
                if delta:
                    yield f"data: {json.dumps({'text': delta})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error("sesiones_resumir_texto LLM error: %s", e)
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream")


LIVE_ANALYSIS_PROMPT = """Estás monitoreando un debate parlamentario en el Congreso del Perú EN VIVO.
Recibirás fragmentos de la transcripción en tiempo real.

Basándote en la transcripción acumulada hasta ahora, genera un análisis breve y actualizado:

**Estado del debate:** [1 línea sobre qué se está debatiendo]
**Posiciones clave:** [qué están diciendo los congresistas, si hay tensiones]
**Puntos de atención:** [algo relevante para un consultor de asuntos públicos]

Sé conciso (máximo 150 palabras). Tono analítico, no descriptivo."""


@app.get("/live/transcribe")
async def live_transcribe(video_id: str = Query(..., description="ID del video de YouTube")):
    """SSE stream que emite líneas de transcripción en tiempo real."""
    api_key = GROQ_API_KEY

    async def generate():
        if not api_key:
            yield f"data: {json.dumps({'error': 'Falta la GROQ_API_KEY'})}\n\n"
            return
        if not video_id:
            yield f"data: {json.dumps({'error': 'Falta el parámetro video_id'})}\n\n"
            return
        try:
            async for item in stream_transcription(video_id, api_key):
                yield f"data: {json.dumps(item, ensure_ascii=False)}\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.post("/live/analyze")
async def live_analyze(request: Request):
    """Analiza el transcript acumulado con Lex. Llamar cada ~60s desde el frontend."""
    body       = await request.json()
    transcript = body.get("transcript", "").strip()
    titulo     = body.get("titulo", "sesión en vivo")
    api_key    = os.getenv("GROQ_API_KEY", "")

    async def generate():
        if not transcript:
            yield f"data: {json.dumps({'error': 'Sin transcripción aún.'})}\n\n"
            return
        # Send only the last 3000 chars to keep tokens low
        excerpt = transcript[-3000:]
        prompt  = f'Sesión: "{titulo}"\n\nTranscripción reciente:\n{excerpt}'
        client  = Groq(api_key=api_key)
        try:
            stream = client.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system", "content": LIVE_ANALYSIS_PROMPT},
                    {"role": "user",   "content": prompt},
                ],
                max_tokens=400,
                temperature=0.3,
                stream=True,
            )
            for chunk in stream:
                delta = chunk.choices[0].delta.content
                if delta:
                    yield f"data: {json.dumps({'text': delta})}\n\n"
            yield "data: [DONE]\n\n"
        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache"})


@app.get("/live", response_class=HTMLResponse)
async def live_page():
    return (Path("static") / "live.html").read_text()


if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", 8732))
    uvicorn.run(app, host="127.0.0.1", port=port)
