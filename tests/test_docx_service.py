"""Tests del conversor Markdown → .docx."""
import io
import zipfile

from docx import Document

from services.docx import markdown_to_docx


def render(md: str) -> Document:
    return Document(io.BytesIO(markdown_to_docx(md)))


def texts(doc) -> list:
    return [p.text for p in doc.paragraphs]


def test_genera_un_docx_valido():
    """Un .docx es un zip con document.xml adentro."""
    data = markdown_to_docx("# Hola")
    assert zipfile.is_zipfile(io.BytesIO(data))
    assert "word/document.xml" in zipfile.ZipFile(io.BytesIO(data)).namelist()


def test_encabezados_por_nivel():
    doc = render("# Uno\n## Dos\n### Tres")
    estilos = [p.style.name for p in doc.paragraphs if p.text in ("Uno", "Dos", "Tres")]
    assert estilos == ["Heading 1", "Heading 2", "Heading 3"]


def test_tabla_markdown_se_convierte_en_tabla_word():
    md = "| Proyecto | Estado |\n|---|---|\n| 14864 | En comisión |\n| 14865 | Archivado |"
    doc = render(md)
    assert len(doc.tables) == 1
    tabla = doc.tables[0]
    assert len(tabla.rows) == 3          # encabezado + 2 filas
    assert tabla.cell(0, 0).text == "Proyecto"
    assert tabla.cell(2, 1).text == "Archivado"


def test_encabezado_de_tabla_en_negrita():
    doc = render("| A | B |\n|---|---|\n| 1 | 2 |")
    run = doc.tables[0].cell(0, 0).paragraphs[0].runs[0]
    assert run.bold


def test_negrita_y_cursiva_inline():
    doc = render("texto **fuerte** y *suave*")
    runs = {r.text: r for p in doc.paragraphs for r in p.runs}
    assert runs["fuerte"].bold
    assert runs["suave"].italic


def test_link_markdown_genera_hipervinculo():
    doc = markdown_to_docx("Ver [el proyecto](https://congreso.gob.pe/pl/14864)")
    xml = zipfile.ZipFile(io.BytesIO(doc)).read("word/document.xml").decode()
    assert "hyperlink" in xml
    assert "el proyecto" in xml


def test_listas_con_vinetas_y_numeradas():
    doc = render("- uno\n- dos\n1. primero")
    estilos = [p.style.name for p in doc.paragraphs if p.text in ("uno", "dos", "primero")]
    assert estilos == ["List Bullet", "List Bullet", "List Number"]


def test_marcado_inline_se_limpia_en_celdas():
    """Las celdas van en texto plano: sin asteriscos sueltos."""
    doc = render("| X |\n|---|\n| **negrita** |")
    assert doc.tables[0].cell(1, 0).text == "negrita"


def test_incluye_pie():
    doc = render("contenido")
    todo = " ".join(texts(doc))
    assert "Generado por Solón" in todo


def test_markdown_vacio_no_revienta():
    assert zipfile.is_zipfile(io.BytesIO(markdown_to_docx("")))
